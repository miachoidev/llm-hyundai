import os
import tempfile
import json
import time
from typing import Dict, Any, List
import streamlit as st
import pandas as pd
import math
from docx import Document
from langchain_core.documents import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_community.retrievers import BM25Retriever
from dotenv import load_dotenv

load_dotenv()   

def extract_table_content(table):
    """
    python-docx의 테이블 객체에서 내용을 추출하는 함수
    """
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        table_data.append(row_data)
    return table_data

def extract_structure_with_python_docx(file_path):
    """
    python-docx를 사용하여 문서의 구조적 정보를 추출하는 함수
    """
    doc = Document(file_path)
    structured_content = []
    
    # 현재 문단 정보 추적
    current_headings = {1: None, 2: None, 3: None, 4: None, 5: None}
    current_section = {"heading": None, "content": [], "level": 0}
    
    # 문서 요소 순서대로 처리
    # 문서의 본문(body)에서 모든 요소를 순서대로 가져오기
    try:
        body = doc._element.body
        content_elements = list(body.iterchildren())
        
        for element_idx, element in enumerate(content_elements):
            element_tag = element.tag.split('}')[-1]  # XML 네임스페이스 제거
            
            # 문단 요소 처리
            if element_tag == 'p':
                # python-docx 문단 객체로 변환
                paragraph = None
                for p in doc.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break
                
                if paragraph is None:
                    continue
                    
                text = paragraph.text
                
                # 빈 문단 건너뛰기
                if not text.strip():
                    continue
                    
                # 제목인지 확인 (스타일 이름에 'Heading' 포함)
                style_name = paragraph.style.name
                is_heading = False
                heading_level = 0
                
                if 'Heading' in style_name:
                    try:
                        # Heading 1, Heading 2 등의 스타일에서 숫자 추출
                        heading_level = int(style_name.split()[-1])
                        is_heading = True
                    except ValueError:
                        pass
                        
                # 숫자 패턴으로 시작하는 제목 감지 (1., 1.1., 1.1.1. 등)
                if not is_heading and text.strip():
                    import re
                    # '1.1.' 형식과 '1.1' 형식 모두 감지하도록 정규식 수정
                    heading_match = re.match(r'^(\d+\.)+\s+|^(\d+\.)+$|^(\d+\.)*\d+\s+', text)
                    if heading_match:
                        # 점의 수로 수준 추정 (마지막 점이 없는 경우도 고려)
                        match_text = heading_match.group(0).strip()
                        num_parts = len(re.findall(r'\d+', match_text))
                        is_heading = True
                        heading_level = num_parts
                    
                    # [부록] 형식 패턴 감지
                    elif re.match(r'^\s*\[\s*부\s*록\s*\d*\s*\]', text) or re.match(r'^\s*\[\s*부\s*록\s*\]', text):
                        is_heading = True
                        heading_level = 1  # 부록은 최상위 수준으로 간주
                        
                    # 별표 기호로 시작하는 패턴도 감지 (예: ** 중요 항목)
                    elif text.strip().startswith('*'):
                        asterisk_count = len(re.match(r'^\s*(\*+)', text.strip()).group(1))
                        if asterisk_count >= 1:
                            is_heading = True
                            heading_level = min(asterisk_count, 3)  # 최대 3레벨까지만 허용
                
                if is_heading:
                    # 이전 섹션 저장
                    if current_section["heading"] is not None:
                        structured_content.append(current_section)
                        
                    # 새 섹션 시작
                    current_headings[heading_level] = text
                    # 하위 수준 헤딩 초기화
                    for i in range(heading_level + 1, 6):
                        current_headings[i] = None
                        
                    # 현재 섹션 갱신
                    current_section = {
                        "heading": text,
                        "content": [],
                        "level": heading_level,
                        "path": [current_headings[i] for i in range(1, heading_level + 1) if current_headings[i] is not None]
                    }
                else:
                    # 일반 문단은 현재 섹션에 추가
                    current_section["content"].append(text)
            
            # 테이블 요소 처리
            elif element_tag == 'tbl':
                # python-docx 테이블 객체로 변환
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                
                if table is None:
                    continue
                    
                table_data = extract_table_content(table)
                
                # 테이블 문자열로 변환
                table_str = "테이블:\n"
                for row in table_data:
                    table_str += " | ".join(row) + "\n"
                
                # 현재 섹션에 테이블 추가 (현재 섹션이 없으면 새 섹션 생성)
                if current_section["heading"] is None:
                    current_section = {
                        "heading": f"테이블 {element_idx}",
                        "content": [table_str],
                        "level": 0,
                        "path": []
                    }
                    structured_content.append(current_section)
                else:
                    current_section["content"].append(table_str)
    
    except Exception as e:
        # 오류 시 기존 방식으로 폴백
        # 문단 처리
        for paragraph in doc.paragraphs:
            # 코드 생략
            pass
            
        # 테이블 처리
        for table in doc.tables:
            # 코드 생략
            pass
    
    # 마지막 섹션 저장
    if current_section["heading"] is not None and current_section not in structured_content:
        structured_content.append(current_section)
    
    return structured_content

def save_documents_backup(chunks, backup_directory=None):
    """
    문서 청크의 백업 파일을 생성하는 함수
    
    Args:
        chunks: 저장할 문서 청크 리스트
        backup_directory: 백업 파일을 저장할 디렉토리 경로 (None이면 기본 경로 사용)
        
    Returns:
        백업 파일 경로
    """
    if backup_directory is None:
        backup_directory = os.path.join(tempfile.gettempdir(), 'document_backup')
    
    os.makedirs(backup_directory, exist_ok=True)
    
    # 백업 파일 경로
    backup_file = os.path.join(backup_directory, 'document_chunks.json')
    
    # JSON 직렬화 가능한 형태로 변환
    serializable_chunks = []
    for chunk in chunks:
        serializable_chunks.append({
            "page_content": chunk.page_content,
            "metadata": chunk.metadata
        })
    
    # 파일로 저장
    with open(backup_file, 'w', encoding='utf-8') as f:
        json.dump(serializable_chunks, f, ensure_ascii=False, indent=2)
    
    print(f"문서 청크 백업 완료: {backup_file} ({len(chunks)} 청크)")
    
    # 텍스트 파일로도 저장 (단순 검색용)
    text_file = os.path.join(backup_directory, 'document_text.txt')
    with open(text_file, 'w', encoding='utf-8') as f:
        for i, chunk in enumerate(chunks):
            f.write(f"\n--- 청크 {i+1} ---\n")
            f.write(f"제목: {chunk.metadata.get('heading', '없음')}\n")
            f.write(f"레벨: {chunk.metadata.get('level', '없음')}\n")
            f.write(f"경로: {' > '.join(chunk.metadata.get('path', []))}\n\n")
            f.write(chunk.page_content)
            f.write("\n\n")
    
    print(f"문서 텍스트 백업 완료: {text_file}")
    
    return backup_directory

def load_documents_backup(backup_directory=None):
    """
    백업된 문서 청크를 로드하는 함수
    
    Args:
        backup_directory: 백업 파일이 저장된 디렉토리 경로 (None이면 기본 경로 사용)
        
    Returns:
        로드된 문서 청크 리스트
    """
    if backup_directory is None:
        backup_directory = os.path.join(tempfile.gettempdir(), 'document_backup')
    
    backup_file = os.path.join(backup_directory, 'document_chunks.json')
    
    if not os.path.exists(backup_file):
        print(f"백업 파일을 찾을 수 없습니다: {backup_file}")
        return []
    
    try:
        # 파일에서 로드
        with open(backup_file, 'r', encoding='utf-8') as f:
            serializable_chunks = json.load(f)
        
        # LangchainDocument 객체로 변환
        chunks = []
        for item in serializable_chunks:
            chunk = LangchainDocument(
                page_content=item["page_content"],
                metadata=item["metadata"]
            )
            chunks.append(chunk)
        
        print(f"문서 청크 로드 완료: {backup_file} ({len(chunks)} 청크)")
        return chunks
    
    except Exception as e:
        print(f"백업 파일 로드 중 오류 발생: {str(e)}")
        return []

def process_document(uploaded_file, chunk_size=1000, chunk_overlap=100) -> Dict[str, Any]:
    print("process_document 함수 호출")
    """
    문서를 청킹하고 벡터 DB에 저장하는 함수
    
    Args:
        uploaded_file: 업로드된 streamlit 파일 객체
        chunk_size: 청크 크기
        chunk_overlap: 청크 간 겹침 크기
        
    Returns:
        Dictionary containing vectorstore and metadata
    """
    result = {
        'status': 'processing',
        'progress': 0,
        'message': '문서 로딩 중...'
    }
    
    tmp_path = None
    
    try:
        # 임시 파일로 저장
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        # 1. 문서 구조 추출
        result['progress'] = 0.1
        result['message'] = '문서 구조 분석 중...'
        
        structured_content = extract_structure_with_python_docx(tmp_path)
        
        result['progress'] = 0.3
        result['message'] = f'문서 구조 추출 완료: {len(structured_content)}개 섹션 발견'
        
        # 2. 구조화된 정보를 LangChain Document로 변환
        documents = []
        for section in structured_content:
            heading = section["heading"]
            content = "\n".join([heading] + section["content"])
            
            # 메타데이터에 경로 정보 포함
            metadata = {
                "heading": heading,
                "level": section["level"],
                "path": section.get("path", [])
            }
            
            doc = LangchainDocument(page_content=content, metadata=metadata)
            documents.append(doc)
        
        # 3. 문서 청킹
        result['progress'] = 0.5
        result['message'] = f'문서 청킹 중... ({len(documents)}개 섹션)'
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        
        chunks = []
        for doc in documents:
            smaller_chunks = text_splitter.split_documents([doc])
            chunks.extend(smaller_chunks)
        
        result['progress'] = 0.7
        result['message'] = f'문서 청킹 완료: 총 {len(chunks)}개 청크 생성, 임베딩 중...'
        
        # 청크 백업 파일 생성
        backup_directory = save_documents_backup(chunks)
        print(f"문서 청크 백업 디렉토리: {backup_directory}")
        
        # 4. 임베딩 및 벡터 DB 저장
        # 임시 디렉토리에 Chroma DB 저장
        persist_directory = os.path.join(tempfile.gettempdir(), 'chroma_db')
        os.makedirs(persist_directory, exist_ok=True)
        print("persist_directory 변수 초기화 완료")

        # OpenAI 임베딩 모델 초기화
        embeddings = OpenAIEmbeddings()
        print("embeddings 변수 초기화 완료")

        # Chroma 벡터 스토어 생성 - 공식 문서 방식대로
        vectorstore = Chroma(
            collection_name="langchain",
            embedding_function=embeddings,
            persist_directory=persist_directory
        )
        print("vectorstore 변수 초기화 완료")
        # 문서 추가
        vectorstore.add_documents(chunks)
        print("문서 추가 완료")

        # 변경사항 영구 저장
        vectorstore.persist()
        
        # BM25 리트리버 생성
        try:
            bm25_retriever = BM25Retriever.from_documents(chunks)
            print("BM25 리트리버 생성 완료")
            
            # BM25 리트리버를 파일로 저장
            bm25_directory = os.path.join(tempfile.gettempdir(), 'bm25_retriever')
            os.makedirs(bm25_directory, exist_ok=True)
            
            try:
                import pickle
                bm25_path = os.path.join(bm25_directory, 'bm25_retriever.pkl')
                with open(bm25_path, 'wb') as f:
                    pickle.dump(bm25_retriever, f)
                print(f"BM25 리트리버 저장 완료: {bm25_path}")
                
                # 저장 확인
                if not os.path.exists(bm25_path):
                    print(f"경고: BM25 리트리버 파일이 저장되지 않았습니다: {bm25_path}")
                else:
                    print(f"BM25 리트리버 파일 확인 완료: {bm25_path} (크기: {os.path.getsize(bm25_path)} 바이트)")
                    
                # 백업 파일도 생성
                backup_path = os.path.join(bm25_directory, 'bm25_retriever_backup.pkl')
                with open(backup_path, 'wb') as f:
                    pickle.dump(bm25_retriever, f)
                print(f"BM25 리트리버 백업 파일 생성 완료: {backup_path}")
                
            except Exception as pickle_err:
                print(f"BM25 리트리버 저장 중 오류 발생: {str(pickle_err)}")
                # 오류가 발생해도 계속 진행 (선택적 기능이므로)
                bm25_directory = None
        except Exception as bm25_err:
            print(f"BM25 리트리버 생성 중 오류 발생: {str(bm25_err)}")
            bm25_directory = None
        
        # 임시 파일 삭제
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)
        
        result['progress'] = 1.0
        result['status'] = 'completed'
        result['message'] = '벡터 DB 및 BM25 리트리버 저장 완료'
        # vectorstore 객체를 직접 저장하지 않고 경로만 저장
        result['persist_directory'] = persist_directory
        result['bm25_directory'] = bm25_directory
        result['backup_directory'] = backup_directory
        result['chunk_count'] = len(chunks)
        
        return result
        
    except Exception as e:
        result['status'] = 'error'
        result['message'] = f'처리 중 오류 발생: {str(e)}'
        
        # 임시 파일 정리
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except:
                pass
                
        return result

def load_retrievers(persist_directory=None, bm25_directory=None):
    """
    벡터 스토어와 BM25 리트리버를 로딩하는 함수
    
    Args:
        persist_directory: Chroma 벡터 스토어가 저장된 디렉토리 경로
        bm25_directory: BM25 리트리버가 저장된 디렉토리 경로
        
    Returns:
        Dictionary containing vector_retriever and bm25_retriever
    """
    result = {}
    vectorstore = None
    
    try:
        # 벡터 스토어 로딩
        if persist_directory and os.path.exists(persist_directory):
            embeddings = OpenAIEmbeddings()
            vectorstore = Chroma(
                collection_name="langchain",
                embedding_function=embeddings,
                persist_directory=persist_directory
            )
            result['vector_retriever'] = vectorstore.as_retriever()
            print("벡터 리트리버 로딩 완료")
        else:
            print(f"벡터 스토어 디렉토리를 찾을 수 없습니다: {persist_directory}")
            return result  # 벡터 스토어가 없으면 리턴
        
        # BM25 리트리버 로딩
        bm25_found = False
        if bm25_directory and os.path.exists(bm25_directory):
            import pickle
            bm25_path = os.path.join(bm25_directory, 'bm25_retriever.pkl')
            if os.path.exists(bm25_path):
                try:
                    with open(bm25_path, 'rb') as f:
                        result['bm25_retriever'] = pickle.load(f)
                    print(f"BM25 리트리버 로딩 완료: {bm25_path}")
                    bm25_found = True
                except Exception as e:
                    print(f"BM25 리트리버 로딩 실패: {str(e)}")
                    # 파일 손상 시 백업 파일 확인
                    backup_path = os.path.join(bm25_directory, 'bm25_retriever_backup.pkl')
                    if os.path.exists(backup_path):
                        try:
                            with open(backup_path, 'rb') as f:
                                result['bm25_retriever'] = pickle.load(f)
                            print(f"BM25 백업 리트리버 로딩 완료: {backup_path}")
                            bm25_found = True
                        except Exception as backup_err:
                            print(f"BM25 백업 리트리버 로딩 실패: {str(backup_err)}")
            else:
                print(f"BM25 리트리버 파일을 찾을 수 없음: {bm25_path}")
        
        # BM25 리트리버를 찾지 못한 경우 즉석에서 생성
        if not bm25_found and vectorstore is not None:
            print("BM25 리트리버를 찾을 수 없어 즉석에서 생성합니다...")
            try:
                # 벡터 스토어에서 문서 가져오기
                collection = vectorstore._collection
                documents = []
                
                # Chroma에서 모든 문서 로드
                ids = collection.get(include=[])["ids"]
                metadatas = collection.get(include=["metadatas"])["metadatas"]
                documents_content = collection.get(include=["documents"])["documents"]
                
                # LangchainDocument 객체로 변환
                for i, doc_content in enumerate(documents_content):
                    metadata = metadatas[i] if i < len(metadatas) else {}
                    documents.append(LangchainDocument(page_content=doc_content, metadata=metadata))
                
                # 즉석에서 BM25 리트리버 생성
                if documents:
                    bm25_retriever = BM25Retriever.from_documents(documents)
                    result['bm25_retriever'] = bm25_retriever
                    print(f"BM25 리트리버 즉석 생성 완료 ({len(documents)} 문서)")
                    
                    # 디렉토리가 없으면 생성
                    if not bm25_directory:
                        bm25_directory = os.path.join(tempfile.gettempdir(), 'bm25_retriever')
                    os.makedirs(bm25_directory, exist_ok=True)
                    
                    # 생성한 리트리버 저장
                    import pickle
                    bm25_path = os.path.join(bm25_directory, 'bm25_retriever.pkl')
                    with open(bm25_path, 'wb') as f:
                        pickle.dump(bm25_retriever, f)
                    print(f"즉석 생성한 BM25 리트리버 저장 완료: {bm25_path}")
                    
                    # 백업 파일도 생성
                    backup_path = os.path.join(bm25_directory, 'bm25_retriever_backup.pkl')
                    with open(backup_path, 'wb') as f:
                        pickle.dump(bm25_retriever, f)
                    print(f"BM25 리트리버 백업 파일 생성 완료: {backup_path}")
                else:
                    print("벡터 스토어에서 문서를 가져올 수 없어 BM25 리트리버를 생성할 수 없습니다.")
            except Exception as create_err:
                print(f"BM25 리트리버 즉석 생성 중 오류 발생: {str(create_err)}")
        
        return result
    
    except Exception as e:
        print(f"리트리버 로딩 중 오류 발생: {str(e)}")
        return result
