import os
import tempfile
import streamlit as st
import pandas as pd
import math
from docx import Document
from langchain_core.documents import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter

def extract_table_content(table):
    """
    python-docx의 테이블 객체에서 내용을 추출하는 함수
    """
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        table_data.append(row_data)
    return table_data

def extract_structure_with_python_docx(file_path):
    """
    python-docx를 사용하여 문서의 구조적 정보를 추출하는 함수
    """
    doc = Document(file_path)
    structured_content = []
    
    # 현재 문단 정보 추적
    current_headings = {1: None, 2: None, 3: None, 4: None, 5: None}
    current_section = {"heading": None, "content": [], "level": 0}
    
    # 문서 요소 순서대로 처리
    # 문서의 본문(body)에서 모든 요소를 순서대로 가져오기
    try:
        body = doc._element.body
        content_elements = list(body.iterchildren())
        print(f"문서 요소 총 {len(content_elements)}개 발견")
        
        for element_idx, element in enumerate(content_elements):
            element_tag = element.tag.split('}')[-1]  # XML 네임스페이스 제거
            
            # 문단 요소 처리
            if element_tag == 'p':
                # python-docx 문단 객체로 변환
                paragraph = None
                for p in doc.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break
                
                if paragraph is None:
                    print(f"경고: 문서 요소 #{element_idx}를 변환할 수 없습니다. 건너뜁니다.")
                    continue
                    
                text = paragraph.text
                print(f"문단 #{element_idx}: {text}")
                
                # 빈 문단 건너뛰기
                if not text.strip():
                    continue
                    
                # 제목인지 확인 (스타일 이름에 'Heading' 포함)
                style_name = paragraph.style.name
                is_heading = False
                heading_level = 0
                
                if 'Heading' in style_name:
                    try:
                        # Heading 1, Heading 2 등의 스타일에서 숫자 추출
                        heading_level = int(style_name.split()[-1])
                        is_heading = True
                        print(f"Heading level: {heading_level}")
                    except ValueError:
                        pass
                        
                # 숫자 패턴으로 시작하는 제목 감지 (1., 1.1., 1.1.1. 등)
                if not is_heading and text.strip():
                    import re
                    # '1.1.' 형식과 '1.1' 형식 모두 감지하도록 정규식 수정
                    heading_match = re.match(r'^(\d+\.)+\s+|^(\d+\.)+$|^(\d+\.)*\d+\s+', text)
                    if heading_match:
                        # 점의 수로 수준 추정 (마지막 점이 없는 경우도 고려)
                        match_text = heading_match.group(0).strip()
                        num_parts = len(re.findall(r'\d+', match_text))
                        is_heading = True
                        heading_level = num_parts
                        print(f"숫자 패턴 제목 감지: '{match_text}', 레벨: {heading_level}")
                    
                    # [부록] 형식 패턴 감지
                    elif re.match(r'^\s*\[\s*부\s*록\s*\d*\s*\]', text) or re.match(r'^\s*\[\s*부\s*록\s*\]', text):
                        is_heading = True
                        heading_level = 1  # 부록은 최상위 수준으로 간주
                        print(f"부록 패턴 제목 감지: '{text}', 레벨: {heading_level}")
                        
                    # 별표 기호로 시작하는 패턴도 감지 (예: ** 중요 항목)
                    elif text.strip().startswith('*'):
                        asterisk_count = len(re.match(r'^\s*(\*+)', text.strip()).group(1))
                        if asterisk_count >= 1:
                            is_heading = True
                            heading_level = min(asterisk_count, 3)  # 최대 3레벨까지만 허용
                            print(f"별표 패턴 제목 감지: '{text}', 레벨: {heading_level}")
                
                if is_heading:
                    # 이전 섹션 저장
                    if current_section["heading"] is not None:
                        structured_content.append(current_section)
                        
                    # 새 섹션 시작
                    current_headings[heading_level] = text
                    # 하위 수준 헤딩 초기화
                    for i in range(heading_level + 1, 6):
                        current_headings[i] = None
                        
                    # 현재 섹션 갱신
                    current_section = {
                        "heading": text,
                        "content": [],
                        "level": heading_level,
                        "path": [current_headings[i] for i in range(1, heading_level + 1) if current_headings[i] is not None]
                    }
                else:
                    # 일반 문단은 현재 섹션에 추가
                    current_section["content"].append(text)
            
            # 테이블 요소 처리
            elif element_tag == 'tbl':
                # python-docx 테이블 객체로 변환
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                
                if table is None:
                    print(f"경고: 테이블 요소 #{element_idx}를 변환할 수 없습니다. 건너뜁니다.")
                    continue
                    
                print(f"테이블 #{element_idx} 처리 중")
                table_data = extract_table_content(table)
                
                # 테이블 문자열로 변환
                table_str = "테이블:\n"
                for row in table_data:
                    table_str += " | ".join(row) + "\n"
                
                # 현재 섹션에 테이블 추가 (현재 섹션이 없으면 새 섹션 생성)
                if current_section["heading"] is None:
                    current_section = {
                        "heading": f"테이블 {element_idx}",
                        "content": [table_str],
                        "level": 0,
                        "path": []
                    }
                    structured_content.append(current_section)
                else:
                    current_section["content"].append(table_str)
    
    except Exception as e:
        print(f"문서 요소 순서 처리 중 오류 발생: {e}")
        print("기본 순서 처리 방식으로 폴백합니다...")
        
        # 오류 시 기존 방식으로 폴백
        # 문단 처리
        for paragraph in doc.paragraphs:
            # (기존 문단 처리 코드...)
            # 코드 생략
            pass
            
        # 테이블 처리
        for table in doc.tables:
            # (기존 테이블 처리 코드...)
            # 코드 생략
            pass
    
    # 마지막 섹션 저장
    if current_section["heading"] is not None and current_section not in structured_content:
        structured_content.append(current_section)
    
    return structured_content

def process_word_document(uploaded_file):
    """
    워드 문서의 구조를 추출하고 청킹
    """
    # 임시 파일로 저장
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name
    
    try:
        # 문서 구조 추출
        structured_content = extract_structure_with_python_docx(tmp_path)
        
        st.info(f"문서 구조 추출 완료: {len(structured_content)}개 섹션 발견")
        
        # 구조화된 정보를 LangChain Document로 변환
        documents = []
        for section in structured_content:
            heading = section["heading"]
            content = "\n".join([heading] + section["content"])
            
            # 메타데이터에 경로 정보 포함
            metadata = {
                "heading": heading,
                "level": section["level"],
                "path": section.get("path", [])
            }
            
            doc = LangchainDocument(page_content=content, metadata=metadata)
            documents.append(doc)
        
        # 필요한 경우 더 작은 청크로 분할
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
        )
        
        chunks = []
        for doc in documents:
            smaller_chunks = text_splitter.split_documents([doc])
            chunks.extend(smaller_chunks)
        
        st.success(f"문서 청킹 완료: 총 {len(chunks)}개 청크 생성")
        
        # 청크 데이터를 표 형식으로 변환
        chunk_data = []
        for i, chunk in enumerate(chunks):
            path = " > ".join(chunk.metadata["path"]) if chunk.metadata["path"] else "-"
            chunk_data.append({
                "번호": i+1,
                "경로": path,
                "내용": chunk.page_content
            })
        
        # 데이터프레임 생성
        chunks_df = pd.DataFrame(chunk_data)
        
        # 페이지네이션 구현
        items_per_page = 20
        total_pages = math.ceil(len(chunk_data) / items_per_page)
        
        col1, col2 = st.columns([4, 1])
        with col2:
            page_number = st.number_input("페이지", min_value=1, max_value=total_pages, value=1)
            
        start_idx = (page_number - 1) * items_per_page
        end_idx = min(start_idx + items_per_page, len(chunk_data))
        
        # 현재 페이지 데이터 표시
        st.subheader(f"청크 목록 (페이지 {page_number}/{total_pages})")
        st.dataframe(
            chunks_df.iloc[start_idx:end_idx],
            use_container_width=True,
            height=600
        )
        
        # 페이지 이동 버튼
        col1, col2 = st.columns(2)
        with col1:
            if page_number > 1:
                if st.button("이전 페이지"):
                    st.session_state.page_number = page_number - 1
                    st.rerun()
        with col2:
            if page_number < total_pages:
                if st.button("다음 페이지"):
                    st.session_state.page_number = page_number + 1
                    st.rerun()
            
        return chunks
    
    finally:
        # 임시 파일 삭제
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

if __name__ == "__main__":
    st.title("워드 문서 구조 분석 및 청킹")
    
    uploaded_file = st.file_uploader("워드 문서(.docx) 업로드", type=["docx"])
    
    if uploaded_file is not None:
        with st.spinner("문서 분석 중..."):
            chunks = process_word_document(uploaded_file)
            
        # 전체 청크 다운로드 기능 추가
        if chunks:
            import json
            
            # 청크를 JSON 형식으로 변환
            chunks_data = []
            for i, chunk in enumerate(chunks):
                chunks_data.append({
                    "id": i+1,
                    "content": chunk.page_content,
                    "metadata": {
                        "heading": chunk.metadata.get("heading", ""),
                        "level": chunk.metadata.get("level", 0),
                        "path": chunk.metadata.get("path", [])
                    }
                })
            
            json_str = json.dumps(chunks_data, ensure_ascii=False, indent=2)
            
            st.download_button(
                label="청킹 결과 다운로드 (JSON)",
                data=json_str,
                file_name="chunked_document.json",
                mime="application/json"
            )