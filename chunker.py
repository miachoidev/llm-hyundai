import os
import tempfile
import json
from typing import Dict, Any, List
from docx import Document
from langchain_core.documents import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter


def extract_table_content(table):
    """
    python-docx의 테이블 객체에서 내용을 추출하는 함수
    """
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        table_data.append(row_data)
    return table_data


def extract_structure_with_python_docx(file_path):
    """
    python-docx를 사용하여 문서의 구조적 정보를 추출하는 함수
    """
    doc = Document(file_path)
    structured_content = []

    # 현재 문단 정보 추적
    current_headings = {1: None, 2: None, 3: None, 4: None, 5: None}
    current_section = {"heading": None, "content": [], "level": 0}

    # 문서 요소 순서대로 처리
    # 문서의 본문(body)에서 모든 요소를 순서대로 가져오기
    try:
        body = doc._element.body
        content_elements = list(body.iterchildren())

        for element_idx, element in enumerate(content_elements):
            element_tag = element.tag.split("}")[-1]  # XML 네임스페이스 제거

            # 문단 요소 처리
            if element_tag == "p":
                # python-docx 문단 객체로 변환
                paragraph = None
                for p in doc.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break

                if paragraph is None:
                    continue

                text = paragraph.text

                # 빈 문단 건너뛰기
                if not text.strip():
                    continue

                # 제목인지 확인 (스타일 이름에 'Heading' 포함)
                style_name = paragraph.style.name
                is_heading = False
                heading_level = 0

                if "Heading" in style_name:
                    try:
                        # Heading 1, Heading 2 등의 스타일에서 숫자 추출
                        heading_level = int(style_name.split()[-1])
                        is_heading = True
                    except ValueError:
                        pass

                # 숫자 패턴으로 시작하는 제목 감지 (1., 1.1., 1.1.1. 등)
                if not is_heading and text.strip():
                    import re

                    # '1.1.' 형식과 '1.1' 형식 모두 감지하도록 정규식 수정
                    heading_match = re.match(
                        r"^(\d+\.)+\s+|^(\d+\.)+$|^(\d+\.)*\d+\s+", text
                    )
                    if heading_match:
                        # 점의 수로 수준 추정 (마지막 점이 없는 경우도 고려)
                        match_text = heading_match.group(0).strip()
                        num_parts = len(re.findall(r"\d+", match_text))
                        is_heading = True
                        heading_level = num_parts

                    # [부록] 형식 패턴 감지
                    elif re.match(r"^\s*\[\s*부\s*록\s*\d*\s*\]", text) or re.match(
                        r"^\s*\[\s*부\s*록\s*\]", text
                    ):
                        is_heading = True
                        heading_level = 1  # 부록은 최상위 수준으로 간주

                    # 별표 기호로 시작하는 패턴도 감지 (예: ** 중요 항목)
                    elif text.strip().startswith("*"):
                        asterisk_count = len(
                            re.match(r"^\s*(\*+)", text.strip()).group(1)
                        )
                        if asterisk_count >= 1:
                            is_heading = True
                            heading_level = min(
                                asterisk_count, 3
                            )  # 최대 3레벨까지만 허용

                if is_heading:
                    # 이전 섹션 저장
                    if current_section["heading"] is not None:
                        structured_content.append(current_section)

                    # 새 섹션 시작
                    current_headings[heading_level] = text
                    # 하위 수준 헤딩 초기화
                    for i in range(heading_level + 1, 6):
                        current_headings[i] = None

                    # 현재 섹션 갱신
                    current_section = {
                        "heading": text,
                        "content": [],
                        "level": heading_level,
                        "path": [
                            current_headings[i]
                            for i in range(1, heading_level + 1)
                            if current_headings[i] is not None
                        ],
                    }
                else:
                    # 일반 문단은 현재 섹션에 추가
                    current_section["content"].append(text)

            # 테이블 요소 처리
            elif element_tag == "tbl":
                # python-docx 테이블 객체로 변환
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break

                if table is None:
                    continue

                table_data = extract_table_content(table)

                # 테이블 문자열로 변환
                table_str = "테이블:\n"
                for row in table_data:
                    table_str += " | ".join(row) + "\n"

                # 현재 섹션에 테이블 추가 (현재 섹션이 없으면 새 섹션 생성)
                if current_section["heading"] is None:
                    current_section = {
                        "heading": f"테이블 {element_idx}",
                        "content": [table_str],
                        "level": 0,
                        "path": [],
                    }
                    structured_content.append(current_section)
                else:
                    current_section["content"].append(table_str)

    except Exception as e:
        # 오류 시 기존 방식으로 폴백
        # 문단 처리
        for paragraph in doc.paragraphs:
            # 코드 생략
            pass

        # 테이블 처리
        for table in doc.tables:
            # 코드 생략
            pass

    # 마지막 섹션 저장
    if (
        current_section["heading"] is not None
        and current_section not in structured_content
    ):
        structured_content.append(current_section)

    return structured_content


def convert_docx_to_chunks(uploaded_file, chunk_size=1000, chunk_overlap=100):
    """
    DOCX 파일을 청크로 변환하는 함수

    Args:
        temp_file_path: 임시 저장된 DOCX 파일 경로
        chunk_size: 청크 크기
        chunk_overlap: 청크 오버랩 크기

    Returns:
        청크 리스트, 저장된 백업 디렉토리 경로
    """
    tmp_path = None
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name
    # 구조적 정보 추출
    structured_content = extract_structure_with_python_docx(tmp_path)

    # Langchain Document 객체로 변환
    documents = []

    for section in structured_content:
        heading = section["heading"]
        content = "\n".join([heading] + section["content"])

        # 메타데이터에 경로 정보 포함
        path_list = section.get("path", [])
        metadata = {
            "heading": heading,
            "level": section["level"],
            "path": " > ".join(path_list) if path_list else "",
        }

        doc = LangchainDocument(page_content=content, metadata=metadata)
        documents.append(doc)

    # 문서를 더 작은 청크로 분할
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )

    # 메타데이터를 보존하며 분할
    chunks = []
    for doc in documents:
        doc_chunks = text_splitter.split_documents([doc])
        chunks.extend(doc_chunks)

    return chunks
